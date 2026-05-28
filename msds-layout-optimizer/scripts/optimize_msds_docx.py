#!/usr/bin/env python3
"""Conservatively apply a bilingual MSDS layout to a DOCX file."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SECTION_RE = re.compile(r"^\s*(?:[1-9]|1[0-6])\.\s+")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 28, start: int = 57, bottom: int = 28, end: int = 57) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def format_run(run, size: float, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph) -> None:
    text = " ".join(paragraph.text.split())
    if not text:
        return

    pf = paragraph.paragraph_format
    if "物质安全资料表" in text or text.upper() in {"MSDS", "MATERIAL SAFETY DATA SHEET"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(10)
        for run in paragraph.runs:
            format_run(run, 26, bold=False)
        return

    if text.startswith("版本") or text.startswith("Version") or "Revision Date" in text or text.startswith("修订日期"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.space_after = Pt(1)
        pf.line_spacing = 1.05
        for run in paragraph.runs:
            format_run(run, 9, color="7F7F7F" if "Revision" in text or "修订" in text else None)
        return

    if SECTION_RE.match(text):
        pf.space_before = Pt(12)
        pf.space_after = Pt(8)
        pf.keep_with_next = True
        for run in paragraph.runs:
            format_run(run, 18, bold=True, color="1F4E79")
        return

    if "Acute toxicity" in text or text.startswith("急性毒性"):
        pf.space_before = Pt(4)
        pf.space_after = Pt(2)
        for run in paragraph.runs:
            format_run(run, 10.5, bold=True)
        return

    pf.space_after = Pt(1)
    pf.line_spacing = 1.05
    for run in paragraph.runs:
        format_run(run, 10.5)


def format_table(table) -> None:
    column_count = len(table.columns)
    dense_table = column_count >= 5

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell, top=28, start=57, bottom=28, end=57)

            if dense_table and row_index == 0:
                set_cell_shading(cell, "D9EAF7")

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                if dense_table:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else paragraph.alignment
                for run in paragraph.runs:
                    size = 8.5 if dense_table else 10.5
                    bold = True if (dense_table and row_index == 0) else None
                    format_run(run, size, bold=bold)


def optimize(input_path: Path, output_path: Path) -> None:
    document = Document(input_path)

    for section in document.sections:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    for paragraph in document.paragraphs:
        format_paragraph(paragraph)

    for table in document.tables:
        format_table(table)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply the Guangzhou Faith-style bilingual MSDS DOCX layout.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    optimize(args.input, args.output)
    print(f"Saved optimized MSDS layout: {args.output}")


if __name__ == "__main__":
    main()
